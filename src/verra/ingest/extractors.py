"""Content extractors for common document formats.

Each extractor returns an ExtractedDocument.  The top-level dispatcher
detect_and_extract() picks the right extractor from the file extension.

Supported formats:
  - PDF  (PyMuPDF / fitz)
  - DOCX (python-docx)
  - TXT / MD (plain text)
  - CSV (csv stdlib)
"""


from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


# Output type


@dataclass
class ExtractedDocument:
    """Normalised output from any extractor."""

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    format: str = "unknown"
    page_count: int = 1


# Code-aware content handling

# Matches fenced code blocks: ```[language]\n...\n```
_FENCED_CODE_RE = re.compile(
    r"```(?P<lang>[a-zA-Z0-9_+-]*)\n(?P<code>.*?)```",
    re.DOTALL,
)

# 4-space-indented blocks (must be preceded by a blank line or start of string)
_INDENTED_CODE_RE = re.compile(
    r"(?:^|\n\n)((?:(?:    |\t)[^\n]*\n?)+)",
)

# HTML <code> tags
_HTML_CODE_RE = re.compile(
    r"<code[^>]*>(?P<code>.*?)</code>",
    re.DOTALL | re.IGNORECASE,
)

# Patterns used to extract meaningful identifiers from code for summarisation
_FUNC_DEF_RE = re.compile(
    r"(?:def|function|func|fn|sub|method)\s+([a-zA-Z_]\w*)",
    re.IGNORECASE,
)
_CLASS_DEF_RE = re.compile(
    r"(?:class|struct|interface|type)\s+([a-zA-Z_]\w*)",
    re.IGNORECASE,
)
_IMPORT_RE = re.compile(
    r"(?:import|require|use|include|from)\s+([a-zA-Z_][\w./]*)",
    re.IGNORECASE,
)
_COMMENT_RE = re.compile(
    r"(?://|#|--)\s*(.+)$",
    re.MULTILINE,
)


def extract_code_blocks(text: str) -> tuple[str, list[dict]]:
    """Separate code blocks from prose in markdown/text content.

    Returns ``(prose_text, code_blocks)`` where ``code_blocks`` is a list of::

        {"language": "python", "code": "...", "context": "surrounding prose"}

    Code blocks are identified by:
    - Markdown fenced blocks (```language ... ```)
    - Indented blocks (4+ spaces, if no fenced blocks overlap)
    - ``<code>`` HTML tags
    """
    code_blocks: list[dict] = []
    # Track ranges already consumed so indented-block pass doesn't double-count.
    consumed_ranges: list[tuple[int, int]] = []

    # --- Fenced blocks ---
    for m in _FENCED_CODE_RE.finditer(text):
        lang = m.group("lang") or ""
        code = m.group("code")
        start, end = m.start(), m.end()
        # Grab up to 100 chars of surrounding prose as context.
        context = text[max(0, start - 100) : start].strip()[-100:]
        code_blocks.append({"language": lang, "code": code, "context": context})
        consumed_ranges.append((start, end))

    # --- HTML <code> tags ---
    for m in _HTML_CODE_RE.finditer(text):
        start, end = m.start(), m.end()
        if any(s <= start < e for s, e in consumed_ranges):
            continue
        code = m.group("code")
        context = text[max(0, start - 100) : start].strip()[-100:]
        code_blocks.append({"language": "", "code": code, "context": context})
        consumed_ranges.append((start, end))

    # --- Indented blocks (only where no fenced/HTML block overlaps) ---
    for m in _INDENTED_CODE_RE.finditer(text):
        start, end = m.start(), m.end()
        if any(s <= start < e for s, e in consumed_ranges):
            continue
        block = m.group(1)
        # Dedent: strip exactly 4 spaces or one tab from each line.
        lines = [
            ln[4:] if ln.startswith("    ") else (ln[1:] if ln.startswith("\t") else ln)
            for ln in block.splitlines()
        ]
        code = "\n".join(lines)
        context = text[max(0, start - 100) : start].strip()[-100:]
        code_blocks.append({"language": "", "code": code, "context": context})
        consumed_ranges.append((start, end))

    # Build prose: remove all identified code spans.
    prose = text
    # Replace in reverse order so offsets remain valid.
    for start, end in sorted(consumed_ranges, reverse=True):
        prose = prose[:start] + prose[end:]

    return prose.strip(), code_blocks


def summarize_code_block(code: str, language: str = "") -> str:
    """Generate a natural language summary of a code block for embedding.

    Extracts function/class names, imports, and inline comments.  This
    produces a short prose description that embeds more meaningfully than
    raw code tokens.
    """
    parts: list[str] = []

    if language:
        parts.append(f"Code ({language}):")

    funcs = _FUNC_DEF_RE.findall(code)
    if funcs:
        parts.append("Defines " + ", ".join(f"`{f}`" for f in funcs[:5]))

    classes = _CLASS_DEF_RE.findall(code)
    if classes:
        parts.append("Classes/types: " + ", ".join(f"`{c}`" for c in classes[:5]))

    imports = _IMPORT_RE.findall(code)
    if imports:
        parts.append("Uses: " + ", ".join(imports[:5]))

    comments = _COMMENT_RE.findall(code)
    meaningful_comments = [c.strip() for c in comments if len(c.strip()) > 5][:3]
    if meaningful_comments:
        parts.append("Comments: " + "; ".join(meaningful_comments))

    if not parts:
        # Fallback: first non-empty line truncated
        first_line = next((ln.strip() for ln in code.splitlines() if ln.strip()), "")
        parts.append(first_line[:120] if first_line else "code block")

    return " ".join(parts)


# Format-specific extractors


def extract_pdf(path: Path) -> ExtractedDocument:
    """Extract text from a PDF.

    Uses pypdf (MIT license) by default.  Falls back to PyMuPDF (AGPL)
    if installed and pypdf is not available — users who need PyMuPDF's
    superior extraction quality can ``pip install pymupdf`` explicitly.
    """
    # Try pypdf first (MIT — safe for any license)
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
        content = "\n\n".join(pages)
        return ExtractedDocument(
            content=content,
            metadata={"file_path": str(path), "file_name": path.name},
            format="pdf",
            page_count=len(reader.pages),
        )
    except ImportError:
        pass

    # Fallback to PyMuPDF if available (AGPL — optional dependency)
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        content = "\n\n".join(p.strip() for p in pages if p.strip())
        return ExtractedDocument(
            content=content,
            metadata={"file_path": str(path), "file_name": path.name},
            format="pdf",
            page_count=len(pages),
        )
    except ImportError:
        pass

    raise ImportError(
        "A PDF library is required: pip install pypdf (recommended, MIT license) "
        "or pip install pymupdf (better quality, AGPL license)"
    )


def extract_docx(path: Path) -> ExtractedDocument:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX extraction: pip install python-docx") from exc

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)

    metadata: dict[str, Any] = {"file_path": str(path), "file_name": path.name}

    if len(content) < 10:
        logger.warning(
            "DOCX extraction produced very little text for %s "
            "(got %d chars) — document may be image-only or corrupt.",
            path,
            len(content),
        )
        metadata["extraction_warning"] = "Document appears empty or image-only"

    return ExtractedDocument(
        content=content,
        metadata=metadata,
        format="docx",
        page_count=1,  # DOCX doesn't have a meaningful page concept here
    )


def extract_text(path: Path) -> ExtractedDocument:
    """Extract plain text (TXT or MD) by reading the file directly."""
    content = path.read_text(encoding="utf-8", errors="replace")
    fmt = "md" if path.suffix.lower() == ".md" else "txt"
    return ExtractedDocument(
        content=content,
        metadata={"file_path": str(path), "file_name": path.name},
        format=fmt,
        page_count=1,
    )


def extract_csv(path: Path) -> ExtractedDocument:
    """Convert a CSV to a readable text representation.

    For large CSVs (>100 rows), splits into logical sections with headers
    repeated so chunks stay self-contained.
    """
    text_rows: list[str] = []
    header: str = ""
    row_count = 0

    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            line = ", ".join(cell.strip() for cell in row if cell.strip())
            if not line:
                continue
            if i == 0:
                header = line
            text_rows.append(line)
            row_count += 1

    # For large CSVs, split into sections with header context
    if row_count > 100:
        sections: list[str] = []
        chunk_size = 50
        data_rows = text_rows[1:]  # skip header
        for start in range(0, len(data_rows), chunk_size):
            batch = data_rows[start : start + chunk_size]
            section = f"{header}\n" + "\n".join(batch)
            section += f"\n\n[Rows {start + 1}-{start + len(batch)} of {len(data_rows)}]"
            sections.append(section)
        content = "\n\n---\n\n".join(sections)
    else:
        content = "\n".join(text_rows)

    return ExtractedDocument(
        content=content,
        metadata={"file_path": str(path), "file_name": path.name, "row_count": row_count},
        format="csv",
        page_count=1,
    )


def extract_xlsx(path: Path) -> ExtractedDocument:
    """Extract text from Excel spreadsheets using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise ValueError(
            f"openpyxl is required for .xlsx files. Install with: pip install openpyxl"
        )

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = []
        header: str = ""
        row_count = 0

        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            line = ", ".join(c for c in cells if c)
            if not line:
                continue
            if i == 0:
                header = line
            rows.append(line)
            row_count += 1

        if rows:
            section = f"## Sheet: {sheet_name}\n\n"
            if row_count > 100:
                # Large sheet — summarize with header + first/last rows
                section += f"{header}\n"
                section += "\n".join(rows[1:51])
                section += f"\n...\n[{row_count} total rows]\n"
                section += "\n".join(rows[-10:])
            else:
                section += "\n".join(rows)
            sections.append(section)

    wb.close()
    content = "\n\n---\n\n".join(sections)

    return ExtractedDocument(
        content=content or "(empty spreadsheet)",
        metadata={
            "file_path": str(path),
            "file_name": path.name,
            "sheet_count": len(wb.sheetnames),
        },
        format="xlsx",
        page_count=len(wb.sheetnames),
    )


# Dispatcher

_EXTENSION_MAP = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_text,
    ".md": extract_text,
    ".markdown": extract_text,
    ".csv": extract_csv,
    ".xlsx": extract_xlsx,
    ".xls": extract_xlsx,
}

SUPPORTED_EXTENSIONS = set(_EXTENSION_MAP.keys())


def detect_and_extract(path: Path) -> ExtractedDocument:
    """Dispatch to the correct extractor based on file extension.

    Raises ValueError for unsupported extensions.
    """
    ext = path.suffix.lower()
    extractor = _EXTENSION_MAP.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file type: {ext!r} (path: {path})")
    return extractor(path)
